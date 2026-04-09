from __future__ import annotations

from datetime import UTC, datetime
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.shared import Inches

from src.report_writer import ReportInsights


def _cell_text(value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, (list, tuple, set, dict)):
        return str(value)
    try:
        if pd.isna(value):
            return ""
    except Exception:
        return str(value)
    return str(value)


def _add_df_table(doc: Document, title: str, df: pd.DataFrame, max_rows: int = 12) -> None:
    doc.add_paragraph(title)
    if df is None or df.empty:
        doc.add_paragraph("No data available.")
        return

    table_df = df.head(max_rows).copy()
    table = doc.add_table(rows=1, cols=len(table_df.columns))
    table.style = "Light List Accent 1"
    hdr = table.rows[0].cells
    for i, c in enumerate(table_df.columns):
        hdr[i].text = str(c)

    for _, r in table_df.iterrows():
        cells = table.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = _cell_text(v)


def _section_status(section_health_df: pd.DataFrame | None, section_name: str) -> str:
    if section_health_df is None or section_health_df.empty:
        return "VALID"
    sub = section_health_df[section_health_df["section_name"] == section_name]
    if sub.empty:
        return "VALID"
    return str(sub.iloc[0].get("section_status", "VALID")).upper()


def _section_reason(section_health_df: pd.DataFrame | None, section_name: str) -> str:
    if section_health_df is None or section_health_df.empty:
        return ""
    sub = section_health_df[section_health_df["section_name"] == section_name]
    if sub.empty:
        return ""
    return str(sub.iloc[0].get("reason", ""))


def _write_section_gate(doc: Document, section_health_df: pd.DataFrame | None, section_name: str) -> bool:
    status = _section_status(section_health_df, section_name)
    if status in {"INVALID", "UNAVAILABLE"}:
        reason = _section_reason(section_health_df, section_name)
        doc.add_paragraph(
            f"Section unavailable for publishable analysis (status={status}). Reason: {reason}"
        )
        return False
    if status == "DEGRADED":
        reason = _section_reason(section_health_df, section_name)
        doc.add_paragraph(
            f"Section is DEGRADED and should be interpreted cautiously. Reason: {reason}"
        )
    return True


def export_word_thesis(
    output_path: Path,
    datasets: dict[str, pd.DataFrame],
    insights: ReportInsights,
    chart_sheets: list[dict[str, Any]],
    metadata: dict[str, str] | None = None,
    run_summary: dict[str, Any] | None = None,
    section_health_df: pd.DataFrame | None = None,
) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    metadata = metadata or {}
    run_summary = run_summary or {}

    doc = Document()

    core_df = datasets.get("Core_Ranking", pd.DataFrame())
    ext_df = datasets.get("Extended_Ranking", pd.DataFrame())
    route_df = datasets.get("Route_Risks", pd.DataFrame())
    choke_df = datasets.get("Chokepoint_Exposure", pd.DataFrame())
    crude_df = datasets.get("Crude_Tracker", pd.DataFrame())
    fuel_df = datasets.get("Fuel_Tracker", pd.DataFrame())
    equity_df = datasets.get("Equity_Tracker", pd.DataFrame())
    factor_df = datasets.get("Factor_Decomposition", pd.DataFrame())
    valuation_df = datasets.get("Valuation", pd.DataFrame())
    scenario_df = datasets.get("Scenario_Analysis", pd.DataFrame())
    analogue_df = datasets.get("Historical_Analogues", pd.DataFrame())
    catalyst_df = datasets.get("Catalyst_Calendar", pd.DataFrame())
    rec_df = datasets.get("Recommendation_Framework", pd.DataFrame())
    writeups_df = datasets.get("Company_Writeups", pd.DataFrame())
    quality_df = datasets.get("Data_Quality", pd.DataFrame())
    assumptions_df = datasets.get("Assumptions", pd.DataFrame())
    source_df = datasets.get("Source_Log", pd.DataFrame())
    missing_df = datasets.get("Missing_Data_Log", pd.DataFrame())
    run_health_df = datasets.get("Run_Health", pd.DataFrame())
    regime_df = datasets.get("Regime_State", pd.DataFrame())
    constraints_df = datasets.get("Market_Constraints", pd.DataFrame())
    constraints_methodology_df = datasets.get("Market_Constraints_Methodology", pd.DataFrame())
    episodes_df = datasets.get("Event_Episodes", pd.DataFrame())
    confidence_df = datasets.get("Confidence_Framework", pd.DataFrame())
    confidence_audit_df = datasets.get("Confidence_Audit", pd.DataFrame())
    peer_df = datasets.get("Peer_Baskets", pd.DataFrame())
    packet_index_df = datasets.get("Company_Case_Packet_Index", pd.DataFrame())
    regression_summary_df = datasets.get("Regression_Summary", pd.DataFrame())
    rolling_beta_summary_df = datasets.get("Rolling_Beta_Summary", pd.DataFrame())
    event_study_summary_df = datasets.get("Event_Study_Summary", pd.DataFrame())

    # 1. Title page
    doc.add_heading("Hormuz Disruption Market Intelligence Memo (Repo 1)", level=0)
    doc.add_paragraph("Upstream market/event/regime intelligence memo generated from Repo 1")
    doc.add_paragraph(datetime.now(UTC).strftime("Generated: %Y-%m-%d %H:%M:%S UTC"))
    doc.add_paragraph(f"Model version: {metadata.get('model_version', '2.0')}")
    doc.add_paragraph(f"Assumptions version: {metadata.get('assumptions_version', '2.0')}")
    doc.add_paragraph(f"Data cut date: {metadata.get('data_cut_date', '')}")
    doc.add_paragraph(f"Run status: {run_summary.get('run_status', insights.run_status)}")
    doc.add_page_break()

    # 2. Executive summary
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(insights.executive_summary)

    # Run health upfront
    doc.add_heading("2. Run-Health Summary", level=1)
    if not run_health_df.empty:
        _add_df_table(doc, "Run Health Checks", run_health_df, max_rows=20)
    else:
        doc.add_paragraph("Run-health table unavailable.")
    if insights.section_caveats:
        doc.add_paragraph("Section caveats:")
        for caveat in insights.section_caveats:
            doc.add_paragraph(f"- {caveat}")

    # 3. Research question
    doc.add_heading("3. Research Question", level=1)
    doc.add_paragraph(insights.research_question)

    # 4. Variant perception / core thesis
    doc.add_heading("4. Variant Perception / Core Thesis", level=1)
    doc.add_paragraph(insights.variant_perception)
    doc.add_paragraph(insights.what_market_missing)

    # 5. Why this matters now
    doc.add_heading("5. Why This Matters Now", level=1)
    doc.add_paragraph(insights.why_now)

    # 6. Universe definition
    doc.add_heading("6. Universe Definition", level=1)
    _add_df_table(
        doc,
        "Core Universe Ranking Snapshot",
        core_df[[c for c in ["core_rank", "company_name", "ticker", "archetype", "final_score", "rating_status"] if c in core_df.columns]],
        max_rows=15,
    )
    _add_df_table(
        doc,
        "Extended Universe Ranking Snapshot",
        ext_df[[c for c in ["extended_rank", "company_name", "ticker", "bucket_classification", "final_score", "rating_status"] if c in ext_df.columns]],
        max_rows=20,
    )

    # 7. Methodology
    doc.add_heading("7. Methodology", level=1)
    doc.add_paragraph(
        "The V2 framework combines route exposure, valuation, oil sensitivity, downstream pass-through, scenario resilience, factor decomposition, catalyst quality, and data confidence penalties. "
        "Core and extended rankings are reported separately. Point estimates are directional and paired with plausible ranges where direct disclosure is unavailable."
    )
    if str(run_summary.get("run_status", "VALID")).upper() != "VALID":
        doc.add_paragraph(
            "Methodology caveat: this run did not fully satisfy publishable data thresholds. Any surviving sections are presented for diagnostic or provisional interpretation."
        )

    # 8. Regime state
    doc.add_heading("8. Regime State", level=1)
    _add_df_table(doc, "Current Regime", regime_df.head(10), max_rows=10)

    # 9. Route exposure analysis
    doc.add_heading("9. Route Exposure Analysis", level=1)
    if _write_section_gate(doc, section_health_df, "route_exposure"):
        _add_df_table(doc, "Chokepoint Exposure", choke_df.head(20), max_rows=20)
        _add_df_table(doc, "Route Risk Commentary", route_df.head(20), max_rows=20)

    # 10. Oil sensitivity / scenario analysis
    doc.add_heading("10. Oil Sensitivity / Scenario Analysis", level=1)
    if _write_section_gate(doc, section_health_df, "market_prices") and _write_section_gate(doc, section_health_df, "fuel_prices"):
        _add_df_table(doc, "Crude Tracker (recent)", crude_df.tail(15), max_rows=15)
        _add_df_table(doc, "Fuel Tracker (recent)", fuel_df.tail(15), max_rows=15)
    if _write_section_gate(doc, section_health_df, "route_exposure"):
        _add_df_table(doc, "Scenario Sensitivities", scenario_df.head(24), max_rows=24)
        _add_df_table(doc, "Event-Path Scenarios", datasets.get("Scenario_Event_Path", pd.DataFrame()).head(24), max_rows=24)

    # 11. Market constraints for downstream valuation
    doc.add_heading("11. Market-to-Valuation Constraints", level=1)
    _add_df_table(doc, "Market Constraint Overlay", constraints_df.head(30), max_rows=30)
    _add_df_table(doc, "Constraint Methodology (Heuristic)", constraints_methodology_df.head(20), max_rows=20)

    # 12. Valuation analysis
    doc.add_heading("12. Valuation Analysis", level=1)
    if _write_section_gate(doc, section_health_df, "valuation"):
        _add_df_table(doc, "Relative Valuation", valuation_df.head(20), max_rows=20)

    # 13. Equity performance / factor decomposition
    doc.add_heading("13. Equity Performance / Factor Decomposition", level=1)
    if _write_section_gate(doc, section_health_df, "equity_performance"):
        _add_df_table(doc, "Equity Tracker (recent)", equity_df.tail(20), max_rows=20)
    if _write_section_gate(doc, section_health_df, "factor_decomposition"):
        _add_df_table(doc, "Factor Decomposition", factor_df.head(20), max_rows=20)

    # 14. Historical analogue analysis
    doc.add_heading("14. Historical Analogue Analysis", level=1)
    if _write_section_gate(doc, section_health_df, "historical_analogues"):
        _add_df_table(doc, "Historical Analogue Summary", analogue_df, max_rows=12)

    # 15. Catalyst map
    doc.add_heading("15. Catalyst Map", level=1)
    if _write_section_gate(doc, section_health_df, "catalyst_map"):
        _add_df_table(doc, "Catalyst Calendar", catalyst_df.head(30), max_rows=30)

    # 16. Event episodes and study
    doc.add_heading("16. Event Episodes and Event Study", level=1)
    _add_df_table(doc, "Event Episodes", episodes_df.head(20), max_rows=20)
    _add_df_table(doc, "Event Study Summary", event_study_summary_df.head(30), max_rows=30)

    # 17. Ranking framework
    doc.add_heading("17. Ranking Framework", level=1)
    _add_df_table(doc, "Core Ranking", core_df.head(15), max_rows=15)
    _add_df_table(doc, "Extended Ranking", ext_df.head(20), max_rows=20)

    # 18. Market-expression screens
    doc.add_heading("18. Market-Expression Screens", level=1)
    if insights.recommendation_allowed:
        doc.add_paragraph("Core top names: " + (", ".join(insights.core_top_names) if insights.core_top_names else "N/A"))
    else:
        doc.add_paragraph("Expression screen suppressed because run-health or ranking publishability thresholds were not met.")

    # 19. Most exposed names
    doc.add_heading("19. Most Exposed Names", level=1)
    doc.add_paragraph("Core weaker names: " + (", ".join(insights.core_bottom_names) if insights.core_bottom_names else "N/A"))

    # 20. Expression framework
    doc.add_heading("20. Upstream Expression Framework", level=1)
    if insights.recommendation_allowed and _section_status(section_health_df, "recommendations") == "VALID":
        doc.add_paragraph(insights.recommendation_text or "No expression summary available.")
        _add_df_table(doc, "Expression Screen Table", rec_df, max_rows=20)
    else:
        doc.add_paragraph("Expression section suppressed: insufficient evidence quality for publishable language.")

    # 21. Risks and limitations
    doc.add_heading("21. Risks and Limitations", level=1)
    for risk in insights.key_risks:
        doc.add_paragraph(f"- {risk}")
    doc.add_paragraph("Potential thesis falsifiers:")
    for f in insights.falsifiers:
        doc.add_paragraph(f"- {f}")

    # 22. Appendix
    doc.add_heading("22. Appendix", level=1)
    _add_df_table(doc, "Company Writeups", writeups_df, max_rows=30)
    _add_df_table(doc, "Data Quality Table", quality_df.head(30), max_rows=30)
    _add_df_table(doc, "Confidence Framework", confidence_df.head(30), max_rows=30)
    _add_df_table(doc, "Confidence Audit", confidence_audit_df.head(30), max_rows=30)
    _add_df_table(doc, "Peer Baskets", peer_df.head(30), max_rows=30)
    _add_df_table(doc, "Company Packet Index", packet_index_df.head(30), max_rows=30)
    _add_df_table(doc, "Regression Summary", regression_summary_df.head(30), max_rows=30)
    _add_df_table(doc, "Rolling Beta Summary", rolling_beta_summary_df.head(30), max_rows=30)

    # 23. Source and estimate legend
    doc.add_heading("23. Source and Estimate Legend", level=1)
    _add_df_table(doc, "Assumptions Registry", assumptions_df.head(30), max_rows=30)
    _add_df_table(doc, "Source Log", source_df.head(40), max_rows=40)
    _add_df_table(doc, "Missing Data Log", missing_df.head(30), max_rows=30)

    doc.add_heading("24. Charts", level=1)
    for meta in chart_sheets:
        doc.add_heading(meta.get("sheet_name", "Chart"), level=2)
        chart_status = str(meta.get("status", "UNKNOWN")).upper()
        doc.add_paragraph(f"Chart status: {chart_status}")
        if meta.get("reason"):
            doc.add_paragraph(f"Reason: {meta.get('reason')}")
        doc.add_paragraph(meta.get("explanation", ""))
        img = Path(meta.get("image_path", ""))
        if img.exists():
            doc.add_picture(str(img), width=Inches(6.4))
        doc.add_paragraph("Takeaway: " + meta.get("takeaway", ""))

    doc.save(output_path)
    return output_path
