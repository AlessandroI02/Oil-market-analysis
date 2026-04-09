from __future__ import annotations

from datetime import UTC, date, datetime, timedelta
from pathlib import Path
import re
from urllib.parse import parse_qs, quote_plus, urlparse
import xml.etree.ElementTree as ET

import pandas as pd
import requests
from docx import Document

from src.assumptions_registry import MissingDataLogger
from src.cache import DiskCache
from src.event_episodes import cluster_event_episodes
from src.fetch_diagnostics import FetchDiagnostics
from src.market_data_providers import fetch_crude_series
from src.source_logger import SourceLogger
from src.storage_paths import ensure_storage_layout, write_csv


_THEME_KEYWORDS = {
    "Hormuz disruption": ["hormuz", "strait", "iran", "gulf shipping"],
    "OPEC / supply": ["opec", "output", "production cut", "quota"],
    "Sanctions": ["sanction", "embargo", "restriction"],
    "Geopolitical attack": ["attack", "drone", "missile", "strike", "houthi"],
    "Inventory surprise": ["inventory", "stockpile", "eia"],
    "Demand shock": ["demand", "china", "consumption", "recession"],
    "Macro / tariff / trade": ["tariff", "trade", "dollar", "rates", "fomc", "inflation"],
    "Refinery disruption": ["refinery", "outage", "maintenance", "shutdown"],
}


_PUBLICATION_TIER_HINTS = {
    "reuters": "Tier 1",
    "associated press": "Tier 1",
    "ap news": "Tier 1",
    "financial times": "Tier 1",
    "wsj": "Tier 1",
    "wall street journal": "Tier 1",
    "bloomberg": "Tier 1",
    "new york times": "Tier 2",
    "washington post": "Tier 2",
    "cnbc": "Tier 2",
    "nikkei": "Tier 2",
    "ft.com": "Tier 1",
    "wsj.com": "Tier 1",
    "bloomberg.com": "Tier 1",
    "axios": "Tier 2",
    "npr": "Tier 2",
    "pbs": "Tier 2",
}


def _extract_domain(url: str) -> str:
    try:
        return urlparse(url).netloc.lower()
    except Exception:
        return ""


def _canonical_url(url: str) -> str:
    try:
        parsed = urlparse(url)
        return f"{parsed.scheme}://{parsed.netloc}{parsed.path}"
    except Exception:
        return url


def _normalize_text(value: str) -> str:
    text = re.sub(r"[\W_]+", " ", str(value).lower())
    return re.sub(r"\s+", " ", text).strip()


def _source_tier(domain: str, publication: str = "") -> str:
    pub = _normalize_text(publication)
    for hint, tier in _PUBLICATION_TIER_HINTS.items():
        if hint in pub:
            return tier

    tier1 = {"eia.gov", "opec.org", "iea.org", "fred.stlouisfed.org", "ec.europa.eu"}
    tier2 = {
        "reuters.com",
        "bloomberg.com",
        "ft.com",
        "wsj.com",
        "apnews.com",
        "cnbc.com",
        "marketwatch.com",
        "nikkei.com",
    }
    tier3 = {
        "seekingalpha.com",
        "investing.com",
        "oilprice.com",
        "finance.yahoo.com",
    }
    if any(d in domain for d in tier1):
        return "Tier 1"
    if any(d in domain for d in tier2):
        return "Tier 2"
    if any(d in domain for d in tier3):
        return "Tier 3"
    return "Tier 4"


def _tier_rank(label: str) -> int:
    return {"Tier 1": 1, "Tier 2": 2, "Tier 3": 3, "Tier 4": 4}.get(str(label), 4)


def _is_wrapper_domain(domain: str) -> bool:
    return str(domain).lower() in {"news.google.com", "google.com", "www.google.com"}


def _fallback_domain_from_publication(publication: str) -> str:
    pub = str(publication).strip().lower()
    if pub.endswith(".com") or pub.endswith(".org") or pub.endswith(".net"):
        return pub
    mapping = {
        "reuters": "reuters.com",
        "bloomberg": "bloomberg.com",
        "wsj": "wsj.com",
        "associated press": "apnews.com",
        "ap news": "apnews.com",
        "financial times": "ft.com",
        "new york times": "nytimes.com",
        "washington post": "washingtonpost.com",
        "cnbc": "cnbc.com",
    }
    for key, domain in mapping.items():
        if key in _normalize_text(publication):
            return domain
    return ""


def _clean_google_link(link: str) -> str:
    try:
        parsed = urlparse(link)
        qs = parse_qs(parsed.query)
        if "url" in qs and qs["url"]:
            return qs["url"][0]
    except Exception:
        pass
    return link


def _resolve_direct_url(url: str, cache: DiskCache) -> str:
    clean_url = str(url or "").strip()
    if clean_url == "":
        return clean_url
    cache_key = f"resolved_link::{clean_url}"
    cached = cache.get_json(cache_key, max_age_hours=72)
    if cached and isinstance(cached, dict) and isinstance(cached.get("url"), str):
        return cached["url"]

    resolved = clean_url
    try:
        if _is_wrapper_domain(_extract_domain(clean_url)):
            response = requests.get(clean_url, timeout=12, allow_redirects=True)
            final_url = str(response.url or "").strip()
            if final_url:
                resolved = final_url
    except Exception:
        resolved = clean_url

    cache.set_json(cache_key, {"url": resolved})
    return resolved


def _theme_from_text(text: str) -> str:
    txt = str(text).lower()
    for theme, keywords in _THEME_KEYWORDS.items():
        if any(k in txt for k in keywords):
            return theme
    return "Other"


def _google_news_search(
    query: str,
    start: date,
    end: date,
    cache: DiskCache,
) -> list[dict[str, object]]:
    cache_key = f"google_news::{query}::{start.isoformat()}::{end.isoformat()}"
    cached = cache.get_json(cache_key, max_age_hours=12)
    if cached and isinstance(cached.get("items"), list):
        return cached["items"]

    q = f"{query} after:{start.isoformat()} before:{end.isoformat()}"
    url = f"https://news.google.com/rss/search?q={quote_plus(q)}&hl=en-US&gl=US&ceid=US:en"
    response = requests.get(url, timeout=20)
    response.raise_for_status()

    root = ET.fromstring(response.text)
    items: list[dict[str, object]] = []
    for item in root.findall("./channel/item"):
        title = (item.findtext("title") or "").strip()
        link = (item.findtext("link") or "").strip()
        pub_date = item.findtext("pubDate")
        source_el = item.find("source")
        publication = source_el.text.strip() if source_el is not None and source_el.text else ""

        clean_link = _clean_google_link(link)
        canonical = _canonical_url(clean_link)
        domain = _extract_domain(clean_link)
        if _is_wrapper_domain(domain):
            domain = _fallback_domain_from_publication(publication) or domain
        source_tier = _source_tier(domain, publication)
        published_dt = pd.to_datetime(pub_date, errors="coerce")

        items.append(
            {
                "headline": title,
                "url": clean_link,
                "canonical_url": canonical,
                "publication": publication or domain,
                "domain": domain,
                "article_date": str(published_dt) if pd.notna(published_dt) else "",
                "source_quality_tier": source_tier,
                "tier_rank": _tier_rank(source_tier),
            }
        )

    cache.set_json(cache_key, {"items": items})
    return items


def _relevance_score(row: pd.Series, move_direction: str) -> float:
    text = f"{row.get('headline', '')} {row.get('publication', '')}".lower()
    score = 0.0
    theme = _theme_from_text(text)
    if theme != "Other":
        score += 2.0

    tier = str(row.get("source_quality_tier", "Tier 4"))
    score += {"Tier 1": 2.4, "Tier 2": 1.8, "Tier 3": 1.0}.get(tier, 0.4)

    if move_direction == "up" and any(k in text for k in ["cut", "attack", "disruption", "sanction", "outage"]):
        score += 1.4
    if move_direction == "down" and any(k in text for k in ["demand", "recession", "oversupply", "slowdown"]):
        score += 1.4
    if any(k in text for k in ["hormuz", "opec", "sanction", "refinery"]):
        score += 0.8
    if _is_wrapper_domain(str(row.get("domain", ""))):
        score -= 0.3
    return score


def _dedupe_articles(df: pd.DataFrame) -> pd.DataFrame:
    if df.empty:
        return df
    out = df.copy()
    out["headline_key"] = out["headline"].astype(str).map(_normalize_text)
    out["domain"] = out.get("domain", pd.Series(dtype=str)).astype(str).str.lower()
    out["source_quality_tier"] = out.get("source_quality_tier", pd.Series(dtype=str)).fillna("Tier 4")
    out["tier_rank"] = out.get("tier_rank", pd.Series(dtype=float)).fillna(out["source_quality_tier"].map(_tier_rank))
    out["article_date"] = pd.to_datetime(out.get("article_date"), errors="coerce")
    out["canonical_url"] = out.get("canonical_url", pd.Series(dtype=str)).astype(str)
    out = out.sort_values(["tier_rank", "article_date"], ascending=[True, False])

    non_wrapper_canonical = ~out["domain"].map(_is_wrapper_domain) & out["canonical_url"].ne("")
    canonical_subset = out[non_wrapper_canonical].drop_duplicates(subset=["canonical_url"], keep="first")
    remaining = out[~non_wrapper_canonical].copy()
    remaining = remaining.drop_duplicates(subset=["headline_key", "publication"], keep="first")
    out = pd.concat([canonical_subset, remaining], ignore_index=True)
    out = out.drop_duplicates(subset=["headline_key"], keep="first")
    return out.drop(columns=["headline_key"], errors="ignore")


def _prefer_high_quality_articles(df: pd.DataFrame) -> pd.DataFrame:
    if df.empty:
        return df
    out = df.copy()
    out["tier_rank"] = out.get("source_quality_tier", pd.Series(dtype=str)).astype(str).map(_tier_rank).fillna(4)
    high_quality_count = int((out["tier_rank"] <= 2).sum())
    if high_quality_count >= 2:
        out = out[out["tier_rank"] <= 3].copy()
    if out.empty:
        return df
    return out


def _build_event_table(brent_daily: pd.DataFrame, wti_daily: pd.DataFrame, threshold_pct: float) -> pd.DataFrame:
    b = brent_daily.copy()
    b["date"] = pd.to_datetime(b["date"], errors="coerce").dt.date
    b = b.sort_values("date")
    b["brent_daily_move_pct"] = pd.to_numeric(b["price"], errors="coerce").pct_change() * 100

    w = wti_daily.copy()
    w["date"] = pd.to_datetime(w["date"], errors="coerce").dt.date
    w = w.sort_values("date")
    w["wti_daily_move_pct"] = pd.to_numeric(w["price"], errors="coerce").pct_change() * 100

    merged = b[["date", "brent_daily_move_pct"]].merge(w[["date", "wti_daily_move_pct"]], on="date", how="left")
    events = merged[merged["brent_daily_move_pct"].abs() >= threshold_pct].copy()
    events["direction"] = events["brent_daily_move_pct"].apply(lambda x: "up" if pd.notna(x) and x >= 0 else "down")
    return events.reset_index(drop=True)


def _build_news_doc(
    output_path: Path,
    episodes_df: pd.DataFrame,
    events_df: pd.DataFrame,
    episode_articles_df: pd.DataFrame,
    catalyst_df: pd.DataFrame,
    threshold_pct: float,
    lookback_days: int,
) -> Path:
    doc = Document()
    doc.add_heading("Oil Market Event Memo (Episode-Based)", level=0)
    doc.add_paragraph(datetime.now(UTC).strftime("Generated: %Y-%m-%d %H:%M:%S UTC"))

    doc.add_heading("1. Methodology", level=1)
    doc.add_paragraph(
        f"Brent daily moves were screened over the trailing {lookback_days} days. Event days are defined as absolute daily Brent moves >= {threshold_pct:.1f}%. "
        "Event days are clustered into episodes to avoid fragmented headline interpretation. Attribution is directional rather than causal proof."
    )

    doc.add_heading("2. Episode Summary", level=1)
    if episodes_df.empty:
        doc.add_paragraph("No episodes were detected in the screened window.")
    else:
        for _, row in episodes_df.iterrows():
            doc.add_paragraph(
                f"- {row['episode_id']} | {row['episode_type']} | {row['start_date']} -> {row['end_date']} | "
                f"Brent cumulative move {float(row['cumulative_brent_move']):.2f}% | confidence {row['attribution_confidence']}"
            )
            if row.get("summary"):
                doc.add_paragraph(str(row["summary"]))
            if row.get("top_articles"):
                doc.add_paragraph(f"Top articles: {row['top_articles']}")

    doc.add_heading("3. Catalyst Theme Totals", level=1)
    if catalyst_df.empty:
        doc.add_paragraph("No catalyst theme totals available.")
    else:
        for _, row in catalyst_df.iterrows():
            doc.add_paragraph(
                f"- {row['theme']}: {int(row['event_count'])} events, avg abs Brent move {float(row['avg_abs_brent_move_pct']):.2f}%"
            )

    doc.add_heading("4. Event-Day Appendix", level=1)
    if events_df.empty:
        doc.add_paragraph("No event days met threshold.")
    else:
        table = doc.add_table(rows=1, cols=7)
        table.style = "Light List Accent 1"
        headers = ["Date", "Brent move %", "WTI move %", "Direction", "Catalyst", "Confidence", "Episode"]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        for _, row in events_df.iterrows():
            cells = table.add_row().cells
            cells[0].text = str(row.get("date", ""))
            cells[1].text = f"{float(row.get('brent_daily_move_pct', 0.0)):.2f}"
            wti_val = row.get("wti_daily_move_pct")
            cells[2].text = "" if pd.isna(wti_val) else f"{float(wti_val):.2f}"
            cells[3].text = str(row.get("direction", ""))
            cells[4].text = str(row.get("candidate_catalyst", ""))
            cells[5].text = str(row.get("attribution_confidence", ""))
            cells[6].text = str(row.get("episode_id", ""))

    doc.add_heading("5. Source Appendix", level=1)
    if episode_articles_df.empty:
        doc.add_paragraph("No sources captured.")
    else:
        src = episode_articles_df.copy()
        src["tier_rank"] = src.get("source_quality_tier", pd.Series(dtype=str)).astype(str).map(_tier_rank).fillna(4)
        src = src.sort_values(["tier_rank", "publication"])
        cols = [c for c in ["publication", "source_quality_tier", "domain", "url", "link_type"] if c in src.columns]
        src = src[cols].drop_duplicates() if cols else pd.DataFrame()
        for _, row in src.iterrows():
            doc.add_paragraph(
                f"- {row.get('publication', '')} ({row.get('source_quality_tier', '')}, {row.get('domain', '')}, {row.get('link_type', '')}): {row.get('url', '')}"
            )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def build_oil_news_report(
    output_word_path: Path,
    debug_dir: Path,
    threshold_pct: float,
    lookback_days: int,
    max_articles_per_event_day: int,
    generate_word: bool = True,
    source_logger: SourceLogger | None = None,
    missing_logger: MissingDataLogger | None = None,
    fetch_diagnostics: FetchDiagnostics | None = None,
) -> dict[str, object]:
    storage = ensure_storage_layout()
    raw_news_dir = storage["data_raw_news"]
    interim_news_dir = storage["data_interim_news"]
    news_cache = DiskCache(storage["data_cache"] / "news_google_rss", rebuild=False)

    end_date = datetime.now(UTC).date()
    start_date = end_date - timedelta(days=lookback_days + 5)

    brent_daily, brent_provider, brent_url, _ = fetch_crude_series("brent", start_date, end_date, diagnostics=fetch_diagnostics)
    wti_daily, wti_provider, wti_url, _ = fetch_crude_series("wti", start_date, end_date, diagnostics=fetch_diagnostics)

    if source_logger:
        if not brent_daily.empty:
            source_logger.add(
                company="GLOBAL",
                field="oil_news_brent_daily",
                source_url=brent_url or "https://finance.yahoo.com/quote/BZ%3DF",
                source_tier="Tier 3",
                evidence_flag="exact",
                comments=f"Provider={brent_provider}",
            )
        if not wti_daily.empty:
            source_logger.add(
                company="GLOBAL",
                field="oil_news_wti_daily",
                source_url=wti_url or "https://finance.yahoo.com/quote/CL%3DF",
                source_tier="Tier 3",
                evidence_flag="exact",
                comments=f"Provider={wti_provider}",
            )

    if brent_daily.empty:
        if missing_logger:
            missing_logger.add(
                company="GLOBAL",
                field_name="oil_news_event_days",
                reason="Brent daily series unavailable for news event detection",
                attempted_sources=[
                    "https://finance.yahoo.com/quote/BZ%3DF",
                    "https://stooq.com/q/d/l/?s=bz.f&i=d",
                    "https://fred.stlouisfed.org/series/DCOILBRENTEU",
                ],
                severity="high",
            )
        events_df = pd.DataFrame(columns=["date", "brent_daily_move_pct", "wti_daily_move_pct", "direction", "candidate_catalyst", "attribution_confidence", "episode_id"])
        articles_df = pd.DataFrame(columns=["event_date", "article_date", "publication", "domain", "headline", "url", "wrapper_url", "canonical_url", "link_type", "theme", "source_quality_tier", "relevance_score", "relevance_note", "episode_id"])
        catalyst_df = pd.DataFrame(columns=["theme", "event_count", "avg_abs_brent_move_pct"])
        episodes_df = pd.DataFrame(columns=["episode_id", "start_date", "end_date", "episode_type", "summary", "involved_event_days", "cumulative_brent_move", "cumulative_wti_move", "mean_abs_move", "peak_move", "attribution_confidence", "top_articles", "episode_source_summary"])
        episode_articles_df = pd.DataFrame(columns=["episode_id"])
    else:
        events_df = _build_event_table(brent_daily, wti_daily, threshold_pct)

        selected_rows: list[dict[str, object]] = []
        raw_candidate_rows: list[dict[str, object]] = []

        for _, event in events_df.iterrows():
            event_date = pd.to_datetime(event["date"]).date()
            query_start = event_date - timedelta(days=1)
            query_end = event_date + timedelta(days=2)
            direction = str(event.get("direction", "up"))

            candidate_articles: list[dict[str, object]] = []
            queries = [
                "brent crude oil",
                "opec oil supply",
                "hormuz shipping",
                "oil sanctions",
                "refinery outage oil",
                "eia oil inventories",
            ]

            for query in queries:
                try:
                    found = _google_news_search(query, query_start, query_end, news_cache)
                    candidate_articles.extend(found)
                    for item in found:
                        raw_candidate_rows.append({"event_date": event_date, "query": query, **item})
                    if fetch_diagnostics:
                        fetch_diagnostics.log_attempt(
                            dataset="oil_news",
                            identifier=f"{event_date}::{query}",
                            provider="google_news_rss",
                            status="success",
                            source_url="https://news.google.com/rss",
                            message=f"articles={len(found)}",
                            row_count=len(found),
                            non_null_count=len(found),
                        )
                except Exception as exc:
                    if fetch_diagnostics:
                        fetch_diagnostics.log_attempt(
                            dataset="oil_news",
                            identifier=f"{event_date}::{query}",
                            provider="google_news_rss",
                            status="failure",
                            source_url="https://news.google.com/rss",
                            message=str(exc),
                            row_count=0,
                            non_null_count=0,
                        )

            candidates_df = pd.DataFrame(candidate_articles)
            candidates_df = _dedupe_articles(candidates_df)
            if candidates_df.empty:
                events_df.loc[events_df["date"] == event_date, "candidate_catalyst"] = "Unattributed"
                events_df.loc[events_df["date"] == event_date, "attribution_confidence"] = "Low"
                continue

            candidates_df["theme"] = candidates_df["headline"].astype(str).apply(_theme_from_text)
            candidates_df["article_date"] = pd.to_datetime(candidates_df["article_date"], errors="coerce")
            candidates_df["relevance_score"] = candidates_df.apply(lambda r: _relevance_score(r, direction), axis=1)
            candidates_df["relevance_note"] = "Theme keyword + source tier + move-direction match; wrapper penalty applied"
            candidates_df = _prefer_high_quality_articles(candidates_df)
            candidates_df["tier_rank"] = candidates_df.get("source_quality_tier", pd.Series(dtype=str)).astype(str).map(_tier_rank).fillna(4)
            selected = (
                candidates_df.sort_values(["tier_rank", "relevance_score", "article_date"], ascending=[True, False, False])
                .drop_duplicates(subset=["publication", "theme"])
                .head(max_articles_per_event_day)
            )
            top_theme = selected["theme"].mode().iloc[0] if not selected.empty else "Unattributed"
            conf_score = selected["relevance_score"].mean() if not selected.empty else 0.0
            top_tier_rank = float(selected["tier_rank"].min()) if not selected.empty else 4.0
            if conf_score >= 4.1 and top_tier_rank <= 2:
                conf = "High"
            elif conf_score >= 3.0:
                conf = "Medium"
            else:
                conf = "Low"

            events_df.loc[events_df["date"] == event_date, "candidate_catalyst"] = top_theme
            events_df.loc[events_df["date"] == event_date, "attribution_confidence"] = conf

            for _, article in selected.iterrows():
                wrapper_url = str(article.get("url", ""))
                resolved_url = _resolve_direct_url(wrapper_url, news_cache)
                resolved_domain = _extract_domain(resolved_url)
                publication = article.get("publication", "")
                if _is_wrapper_domain(resolved_domain):
                    resolved_domain = _fallback_domain_from_publication(str(publication)) or str(article.get("domain", ""))
                source_tier = _source_tier(resolved_domain, str(publication))
                link_type = "direct" if not _is_wrapper_domain(_extract_domain(resolved_url)) else "wrapper"
                canonical_url = _canonical_url(resolved_url)

                selected_rows.append(
                    {
                        "event_date": event_date,
                        "article_date": pd.to_datetime(article.get("article_date"), errors="coerce").date() if pd.notna(article.get("article_date")) else None,
                        "publication": publication,
                        "domain": resolved_domain,
                        "headline": article.get("headline", ""),
                        "url": resolved_url,
                        "wrapper_url": wrapper_url,
                        "canonical_url": canonical_url,
                        "link_type": link_type,
                        "theme": article.get("theme", "Other"),
                        "source_quality_tier": source_tier,
                        "relevance_score": article.get("relevance_score", 0.0),
                        "relevance_note": article.get("relevance_note", ""),
                    }
                )

        articles_df = pd.DataFrame(selected_rows)
        events_df["candidate_catalyst"] = events_df.get("candidate_catalyst", "Unattributed").fillna("Unattributed")
        events_df["attribution_confidence"] = events_df.get("attribution_confidence", "Low").fillna("Low")

        if not events_df.empty:
            catalyst_df = (
                events_df.assign(abs_brent_move=lambda d: d["brent_daily_move_pct"].abs())
                .groupby("candidate_catalyst", as_index=False)
                .agg(event_count=("date", "count"), avg_abs_brent_move_pct=("abs_brent_move", "mean"))
                .rename(columns={"candidate_catalyst": "theme"})
                .sort_values(["event_count", "avg_abs_brent_move_pct"], ascending=[False, False])
            )
        else:
            catalyst_df = pd.DataFrame(columns=["theme", "event_count", "avg_abs_brent_move_pct"])

        episodes_df, episode_articles_df = cluster_event_episodes(events_df, articles_df)
        if not episode_articles_df.empty and "episode_id" in episode_articles_df.columns:
            articles_df = episode_articles_df.copy()
        if not episodes_df.empty:
            date_map: dict[str, str] = {}
            for _, ep in episodes_df.iterrows():
                for part in str(ep.get("involved_event_days", "")).split(","):
                    key = part.strip()
                    if key:
                        date_map[key] = str(ep.get("episode_id", ""))
            events_df["episode_id"] = events_df["date"].astype(str).map(date_map).fillna("")
        else:
            events_df["episode_id"] = ""

        raw_candidates_df = pd.DataFrame(raw_candidate_rows)
        if not raw_candidates_df.empty:
            write_csv(raw_candidates_df, raw_news_dir / "oil_news_raw_candidates.csv")
            write_csv(raw_candidates_df, interim_news_dir / "oil_news_raw_candidates.csv")

    debug_dir.mkdir(parents=True, exist_ok=True)
    event_path = debug_dir / "oil_news_event_days.csv"
    articles_path = debug_dir / "oil_news_articles.csv"
    catalysts_path = debug_dir / "oil_news_catalysts.csv"
    episodes_path = debug_dir / "event_episodes.csv"
    episode_articles_path = debug_dir / "event_episode_articles.csv"
    episode_summary_path = debug_dir / "oil_news_episode_summary.csv"

    events_df.to_csv(event_path, index=False)
    articles_df.to_csv(articles_path, index=False)
    catalyst_df.to_csv(catalysts_path, index=False)
    episodes_df.to_csv(episodes_path, index=False)
    episode_articles_df.to_csv(episode_articles_path, index=False)
    episodes_df.to_csv(episode_summary_path, index=False)

    write_csv(events_df, interim_news_dir / "oil_news_event_days.csv")
    write_csv(articles_df, interim_news_dir / "oil_news_articles.csv")
    write_csv(catalyst_df, interim_news_dir / "oil_news_catalysts.csv")
    write_csv(episodes_df, interim_news_dir / "event_episodes.csv")

    news_doc_path = None
    if generate_word:
        news_doc_path = _build_news_doc(
            output_path=output_word_path,
            episodes_df=episodes_df,
            events_df=events_df,
            episode_articles_df=episode_articles_df if not episode_articles_df.empty else articles_df,
            catalyst_df=catalyst_df,
            threshold_pct=threshold_pct,
            lookback_days=lookback_days,
        )

    return {
        "word_path": news_doc_path,
        "oil_news_event_days": event_path,
        "oil_news_articles": articles_path,
        "oil_news_catalysts": catalysts_path,
        "oil_news_episode_summary": episode_summary_path,
        "event_episodes": episodes_path,
        "event_episode_articles": episode_articles_path,
        "events_df": events_df,
        "articles_df": articles_df,
        "catalysts_df": catalyst_df,
        "episodes_df": episodes_df,
        "episode_articles_df": episode_articles_df,
    }
